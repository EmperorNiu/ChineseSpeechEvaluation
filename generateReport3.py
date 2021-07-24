from docx import Document
from docx import table
import sys
from sqlalchemy import Column,String,Integer,ForeignKey,create_engine,PrimaryKeyConstraint
from sqlalchemy.orm import sessionmaker
from sqlalchemy.ext.declarative import declarative_base
from collections import Counter

student_id = sys.argv[1]
homework_doc_id = sys.argv[2]
student_name = sys.argv[3]
teacher = sys.argv[4]
title = sys.argv[5]
time = sys.argv[6]
audio = sys.argv[7]
doc_path = './report_template2.docx'

# student_id = 'PS001'
# homework_doc_id = 69

Base = declarative_base()
class HomeworkDoc(Base):
    __tablename__ = 'homework_doc'
    homework_doc_id = Column(Integer,primary_key=True) 
    title = Column(String)
    describe = Column(String)
    position = Column(String)
class HomeworkResult(Base):
    __tablename__ = 'student_homework_result'
    student_homework_result_id = Column(Integer,primary_key=True)
    homework_doc_id_refer = Column(Integer)
    student_id_refer = Column(String)
    tone_accuracy = Column(Integer)
    intonation_accuracy = Column(Integer)
    fluency = Column(Integer)
    comment = Column(String)

engine = create_engine('mysql+mysqlconnector://root:123456@localhost:3306/charword')
DBSession = sessionmaker(bind=engine)
db = DBSession()

re = db.query(HomeworkResult).filter_by(homework_doc_id_refer=homework_doc_id,student_id_refer=student_id).all()[0]
hdoc = db.query(HomeworkDoc).filter_by(homework_doc_id=homework_doc_id).all()[0]


# doc = Document(hdoc.position)
doc = Document(doc_path)

s1 = int(re.tone_accuracy*0.2)
s2 = int(re.intonation_accuracy*0.1)
s3 = int(re.fluency*0.1)

doc.tables[0].rows[7].cells[9].text = str(s1)
doc.tables[0].rows[7].cells[10].text = str(s2)
doc.tables[0].rows[7].cells[12].text = str(s3)

doc.tables[0].rows[8].cells[0].text = '评语：\n' + re.comment


doc.tables[0].rows[0].cells[6].text = title
doc.tables[0].rows[2].cells[6].text = student_name
doc.tables[0].rows[3].cells[6].text = teacher
doc.tables[0].rows[2].cells[13].text = student_id
doc.tables[0].rows[1].cells[6].text = time
doc.tables[0].rows[1].cells[13].text = time
doc.tables[0].rows[0].cells[13].text = audio
score = s1 + s2 + s3
doc.tables[0].rows[3].cells[13].text = str(int(score))


dname = './download/doc/report_' + student_id + '_' + str(homework_doc_id) + '.docx'
doc.save(dname)
# print("success")
db.close()
# tmp_name = './report_' + student_id + '_' + str(homework_doc_id) + '.docx'
# doc.save(tmp_name)
